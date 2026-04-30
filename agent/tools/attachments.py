"""read_attachment — pull (and lazily extract) text from a chat attachment.

Realtors drop files into the prompt box. /api/ai/attachments persists each
upload as an Attachment row. /api/ai/task hydrates the rows the realtor
referenced in this turn and forwards them to the chat_turn endpoint, which
appends a `[Attached <filename> — id <id>; call read_attachment ...]`
marker to the user message so the model knows to call this tool.

Extraction rules:
  - Images are skipped on upload — the model already gets them via the
    image_url path. We just return the public_url and a hint.
  - PDF / DOCX / XLSX / text-ish files are extracted on demand the first
    time the agent asks. We persist the extracted text so subsequent reads
    are cheap.
  - Anything else returns a polite error the agent can surface to the user.

Failures NEVER throw out of the tool. The agent gets {"error": "..."} and
decides what to tell the user.
"""

from __future__ import annotations

import io
import json
from typing import Any

import httpx
from agents import RunContextWrapper, function_tool

from db import supabase
from security.context import AgentContext

# Cap the text we hand back to the model. 50k chars is roughly 12k tokens —
# well above what the model can usefully consume per attachment but enough
# to cover a typical disclosure packet or rent roll.
_MAX_TEXT_CHARS = 50_000
# Cap the bytes we'll pull from Supabase for extraction. Matches the upload
# route's 25 MB ceiling.
_MAX_DOWNLOAD_BYTES = 25 * 1024 * 1024
# httpx timeout — the public URL is on the same Supabase project so this is
# generous on purpose.
_DOWNLOAD_TIMEOUT = 15.0

_TEXTLIKE_MIMES = {
    "text/plain",
    "text/csv",
    "text/markdown",
    "application/json",
}


def _truncate(text: str, limit: int = _MAX_TEXT_CHARS) -> str:
    if len(text) <= limit:
        return text
    return text[: limit - 1] + "…"


async def _download(public_url: str) -> bytes:
    """Fetch the attachment bytes from its public URL. Caps at _MAX_DOWNLOAD_BYTES."""
    async with httpx.AsyncClient(timeout=_DOWNLOAD_TIMEOUT) as client:
        resp = await client.get(public_url)
        resp.raise_for_status()
        # httpx already loads the full body — clip if oversized so we don't
        # blow up extraction on a runaway file.
        data = resp.content
        if len(data) > _MAX_DOWNLOAD_BYTES:
            data = data[:_MAX_DOWNLOAD_BYTES]
        return data


def _extract_pdf(data: bytes) -> str:
    """PDF text extraction via pypdf. Returns '' on parse failure."""
    try:
        from pypdf import PdfReader  # type: ignore[import-not-found]
    except ImportError:
        return "[extraction unavailable: pypdf not installed in this environment]"
    try:
        reader = PdfReader(io.BytesIO(data))
        chunks: list[str] = []
        for page in reader.pages:
            try:
                chunks.append(page.extract_text() or "")
            except Exception:
                # One bad page shouldn't kill the whole document.
                continue
            # Bail early once we've collected enough — no point parsing a
            # 500-page disclosure when we'll truncate to 50k chars anyway.
            if sum(len(c) for c in chunks) > _MAX_TEXT_CHARS:
                break
        return "\n".join(chunks).strip()
    except Exception as exc:
        return f"[pdf parse failed: {exc}]"


def _extract_docx(data: bytes) -> str:
    """DOCX text extraction via python-docx."""
    try:
        from docx import Document  # type: ignore[import-not-found]
    except ImportError:
        return "[extraction unavailable: python-docx not installed in this environment]"
    try:
        doc = Document(io.BytesIO(data))
        paragraphs = [p.text for p in doc.paragraphs if p.text]
        # Tables sometimes carry the meat of a real-estate doc (rent rolls,
        # comparison sheets) so flatten them after the body paragraphs.
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
                if cells:
                    paragraphs.append(" | ".join(cells))
        return "\n".join(paragraphs).strip()
    except Exception as exc:
        return f"[docx parse failed: {exc}]"


def _extract_xlsx(data: bytes) -> str:
    """XLSX text extraction via openpyxl. Flattens to TSV-ish text."""
    try:
        from openpyxl import load_workbook  # type: ignore[import-not-found]
    except ImportError:
        return "[extraction unavailable: openpyxl not installed in this environment]"
    try:
        wb = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
        out: list[str] = []
        for sheet in wb.worksheets:
            out.append(f"# Sheet: {sheet.title}")
            for row in sheet.iter_rows(values_only=True):
                cells = ["" if v is None else str(v) for v in row]
                if any(c for c in cells):
                    out.append("\t".join(cells))
                if sum(len(line) for line in out) > _MAX_TEXT_CHARS:
                    break
            if sum(len(line) for line in out) > _MAX_TEXT_CHARS:
                break
        wb.close()
        return "\n".join(out).strip()
    except Exception as exc:
        return f"[xlsx parse failed: {exc}]"


def _extract_text(data: bytes, mime: str) -> str:
    """Plain text / CSV / markdown / JSON. Decode best-effort, pretty JSON."""
    try:
        text = data.decode("utf-8", errors="replace")
    except Exception as exc:
        return f"[decode failed: {exc}]"
    if mime == "application/json":
        # Pretty-print so the model can scan structure; fall back to raw
        # if it isn't actually JSON.
        try:
            parsed = json.loads(text)
            return json.dumps(parsed, indent=2, default=str)
        except (json.JSONDecodeError, ValueError):
            return text
    return text


async def _persist_extraction(
    db: Any, attachment_id: str, space_id: str, text: str, status: str
) -> None:
    """Best-effort write-back so subsequent reads skip extraction."""
    try:
        await (
            db.table("Attachment")
            .update({"extractedText": text, "extractionStatus": status})
            .eq("id", attachment_id)
            .eq("spaceId", space_id)
            .execute()
        )
    except Exception:
        # Persisting is a cache, not a contract — swallow.
        pass


@function_tool
async def read_attachment(
    ctx: RunContextWrapper[AgentContext],
    attachment_id: str,
) -> dict[str, Any]:
    """Read the contents of a chat attachment by id.

    The user message footer lists attachment ids as `[Attached <filename> — id <id>]`.
    For text-bearing files (PDF, DOCX, XLSX, txt/csv/md/json), this returns
    `{filename, mime_type, text}`. For images, the model already received the
    image via vision, so this returns `{filename, mime_type, text: null,
    public_url, hint}`. On failure, returns `{error: "..."}`.
    """
    space_id = ctx.context.space_id
    db = await supabase()

    try:
        result = await (
            db.table("Attachment")
            .select(
                'id,filename,"mimeType","publicUrl","extractedText","extractionStatus"'
            )
            .eq("id", attachment_id)
            .eq("spaceId", space_id)  # tenant boundary — never skip
            .single()
            .execute()
        )
    except Exception as exc:
        return {"error": f"attachment lookup failed: {exc}"}

    row = result.data
    if not row:
        return {"error": "attachment not found in this workspace"}

    filename = row.get("filename") or "attachment"
    mime_type = row.get("mimeType") or "application/octet-stream"
    public_url = row.get("publicUrl") or ""
    status = row.get("extractionStatus") or "pending"
    cached = row.get("extractedText")

    # ---- already extracted ----
    if status == "done" and cached:
        return {
            "filename": filename,
            "mime_type": mime_type,
            "text": _truncate(cached),
        }

    # ---- image: model has it via image_url, just point at it ----
    if status == "skipped" or mime_type.startswith("image/"):
        return {
            "filename": filename,
            "mime_type": mime_type,
            "text": None,
            "public_url": public_url,
            "hint": "image — described in the message context as image_url",
        }

    # ---- previously failed: don't loop on the same broken file ----
    if status == "failed":
        return {
            "filename": filename,
            "mime_type": mime_type,
            "text": None,
            "error": "extraction previously failed for this attachment",
        }

    # ---- pending: extract now ----
    if not public_url:
        await _persist_extraction(db, attachment_id, space_id, "", "failed")
        return {"error": "attachment has no public url"}

    try:
        data = await _download(public_url)
    except Exception as exc:
        await _persist_extraction(db, attachment_id, space_id, "", "failed")
        return {"error": f"download failed: {exc}"}

    text: str
    if mime_type == "application/pdf":
        text = _extract_pdf(data)
    elif (
        mime_type
        == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    ):
        text = _extract_docx(data)
    elif (
        mime_type
        == "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    ):
        text = _extract_xlsx(data)
    elif mime_type in _TEXTLIKE_MIMES:
        text = _extract_text(data, mime_type)
    else:
        await _persist_extraction(db, attachment_id, space_id, "", "failed")
        return {
            "filename": filename,
            "mime_type": mime_type,
            "text": None,
            "error": "unsupported mime type",
        }

    text = (text or "").strip()
    if not text:
        await _persist_extraction(db, attachment_id, space_id, "", "failed")
        return {
            "filename": filename,
            "mime_type": mime_type,
            "text": None,
            "error": "extraction produced no text",
        }

    truncated = _truncate(text)
    # Cache the truncated form — it's the only thing we'll ever serve.
    await _persist_extraction(db, attachment_id, space_id, truncated, "done")

    return {
        "filename": filename,
        "mime_type": mime_type,
        "text": truncated,
    }
